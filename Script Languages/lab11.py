import re
import docx
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image


# Load the book
with open("C:\\Users\\Khalid\\Desktop\\Joking apart.txt", "r", encoding='utf-8') as file:
    book_text = file.read()

find_title = re.search('Title: (.*)\n', book_text)
find_author = re.search('Author: (.*)\n', book_text)

if find_title and find_author:
    title = find_title.group(1)
    author = find_author.group(1)


first_chapter_start = book_text.find('Just to show the sort of')
first_chapter_end = book_text.find('have seen at one time or another.')
first_chapter = book_text[first_chapter_start:first_chapter_end]

# Evaluate the number of words in each paragraph of the first chapter and create a plot
paragraphs = first_chapter.split('\n\n')
paragraph_lengths = [len(paragraph.split()) for paragraph in paragraphs]

# PART 4: Create a plot showing a distribution of lengths of paragraphs in the first chapter
plt.hist(paragraph_lengths, bins=20)
plt.xlabel('Length of paragraphs')
plt.ylabel('Number of paragraphs')
plt.title('Distribution of Paragraph Lengths in First Chapter')
plt.grid(True)
plt.savefig('plot.png')
plt.show()

# Load a picture related to the book content and resize it
image = Image.open("C:\\Users\\Khalid\\PycharmProjects\\Lab11\\image1.jpg")
resized_image = image.resize((700, 700))

# Load a black-and-white logo, rotate it, and paste it into the first picture
logo = Image.open("C:\\Users\\Khalid\\PycharmProjects\\Lab11\\logo.jpg")
rotated_logo = logo.rotate(55)
resized_image.paste(rotated_logo, (0, 0))
resized_image.save('resized_image.jpg')

# PART 8: Create a Word document
doc = docx.Document()
doc.add_heading(title, 0)
doc.add_picture('resized_image.jpg', width=docx.shared.Inches(6))
doc.add_heading('Author: ' + author, level=1)
doc.add_heading('Khalid Muzaffar', level=2)
doc.add_page_break()

# Add the plot to the Word document
doc.add_picture('plot.png', width=docx.shared.Inches(7))

# Add a text paragraph including description of the picture content
text_parts = [
    ("The plot above shows the distribution of paragraph lengths in the first chapter. The data is as follows:\n", False, '000000', 12),
    (f"- Number of paragraphs: {len(paragraphs)}\n", True, 'FF0000', 14),
    (f"- Number of words in the first chapter: {sum(paragraph_lengths)}\n", True, 'FF0000', 14),
    (f"- Minimal number of words in paragraph: {min(paragraph_lengths)}\n", True, 'FF0000', 14),
    (f"- Maximal number of words in paragraph: {max(paragraph_lengths)}\n", True, 'FF0000', 14),
    (f"- Average number of words in paragraph: {np.mean(paragraph_lengths):.2f}\n", True, 'FF0000', 14),
    ("The source of the data is the first chapter of the book.", False, '000000', 12)
]

# Format the text paragraph
paragraph = doc.add_paragraph()
for text, bold, color, size in text_parts:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = docx.shared.Pt(size)
    run.font.color.rgb = docx.shared.RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))

doc.save("report.docx")
